"""
Utilities for generating PDF and DOCX documents from report content
"""

from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from docx.shared import Inches

def generate_pdf(content: str) -> bytes:
    """
    Generate a PDF document from the report content
    
    Args:
        content (str): The report content to convert to PDF
        
    Returns:
        bytes: The PDF document as bytes
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Create custom style for the content
    content_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        leading=14,
        spaceAfter=12
    )
    
    # Build the document
    elements = []
    
    # Add title
    elements.append(Paragraph("Research Report", styles['Title']))
    elements.append(Spacer(1, 12))
    
    # Add content paragraphs
    for paragraph in content.split('\n\n'):
        if paragraph.strip():
            elements.append(Paragraph(paragraph, content_style))
            elements.append(Spacer(1, 12))
    
    doc.build(elements)
    
    # Get the value of the BytesIO buffer
    pdf = buffer.getvalue()
    buffer.close()
    
    return pdf

def generate_docx(content: str) -> bytes:
    """
    Generate a DOCX document from the report content
    
    Args:
        content (str): The report content to convert to DOCX
        
    Returns:
        bytes: The DOCX document as bytes
    """
    doc = Document()
    
    # Add title
    doc.add_heading('Research Report', 0)
    
    # Add content paragraphs
    for paragraph in content.split('\n\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    
    # Get the value of the BytesIO buffer
    docx = buffer.getvalue()
    buffer.close()
    
    return docx